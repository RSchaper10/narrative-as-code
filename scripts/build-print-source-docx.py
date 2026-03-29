#!/usr/bin/env python3

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build a print-source DOCX from the compiled Markdown manuscript."
    )
    parser.add_argument(
        "--input",
        default="build/manuscript-draft.md",
        help="Path to the compiled Markdown manuscript.",
    )
    parser.add_argument(
        "--output",
        default="build/manuscript-print-source.docx",
        help="Path to the output DOCX file.",
    )
    return parser.parse_args()


def split_blocks(text: str) -> list[list[str]]:
    blocks: list[list[str]] = []
    current: list[str] = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.strip():
            current.append(line)
            continue
        if current:
            blocks.append(current)
            current = []

    if current:
        blocks.append(current)

    return blocks


def get_or_add_paragraph_style(document: Document, name: str):
    styles = document.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.15

    title = get_or_add_paragraph_style(document, "PrintSourceTitle")
    title.base_style = normal
    title.font.name = "Times New Roman"
    title.font.size = Pt(20)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)

    byline = get_or_add_paragraph_style(document, "PrintSourceByline")
    byline.base_style = normal
    byline.font.name = "Times New Roman"
    byline.font.size = Pt(12)
    byline.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.paragraph_format.space_after = Pt(12)

    imprint = get_or_add_paragraph_style(document, "PrintSourceImprint")
    imprint.base_style = normal
    imprint.font.name = "Times New Roman"
    imprint.font.size = Pt(11)
    imprint.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imprint.paragraph_format.space_after = Pt(12)

    chapter = get_or_add_paragraph_style(document, "PrintSourceChapterHeading")
    chapter.base_style = normal
    chapter.font.name = "Times New Roman"
    chapter.font.size = Pt(16)
    chapter.font.bold = True
    chapter.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter.paragraph_format.space_after = Pt(18)

    subheading = get_or_add_paragraph_style(document, "PrintSourceSubheading")
    subheading.base_style = normal
    subheading.font.name = "Times New Roman"
    subheading.font.size = Pt(12)
    subheading.font.bold = True
    subheading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheading.paragraph_format.space_after = Pt(12)

    copyright_style = get_or_add_paragraph_style(document, "PrintSourceCopyright")
    copyright_style.base_style = normal
    copyright_style.font.name = "Times New Roman"
    copyright_style.font.size = Pt(10)
    copyright_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    copyright_style.paragraph_format.space_after = Pt(6)
    copyright_style.paragraph_format.line_spacing = 1.1

    dedication = get_or_add_paragraph_style(document, "PrintSourceDedication")
    dedication.base_style = normal
    dedication.font.name = "Times New Roman"
    dedication.font.size = Pt(11)
    dedication.font.italic = True
    dedication.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dedication.paragraph_format.space_after = Pt(12)

    body = get_or_add_paragraph_style(document, "PrintSourceBody")
    body.base_style = normal
    body.font.name = "Times New Roman"
    body.font.size = Pt(11)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body.paragraph_format.first_line_indent = Inches(0.25)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.line_spacing = 1.15

    body_first = get_or_add_paragraph_style(document, "PrintSourceBodyFirst")
    body_first.base_style = normal
    body_first.font.name = "Times New Roman"
    body_first.font.size = Pt(11)
    body_first.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body_first.paragraph_format.first_line_indent = Inches(0)
    body_first.paragraph_format.space_after = Pt(0)
    body_first.paragraph_format.line_spacing = 1.15

    scene_break = get_or_add_paragraph_style(document, "PrintSourceSceneBreak")
    scene_break.base_style = normal
    scene_break.font.name = "Times New Roman"
    scene_break.font.size = Pt(11)
    scene_break.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scene_break.paragraph_format.space_after = Pt(6)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    section.gutter = Inches(0)


def add_markdown_runs(paragraph, text: str) -> None:
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        run_text = part
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run.bold = True
            run_text = part[2:-2]
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            run.italic = True
            run_text = part[1:-1]
        run.text = run_text


def add_markdown_paragraph(document: Document, text: str, style: str) -> None:
    paragraph = document.add_paragraph(style=style)
    for idx, line in enumerate(text.split("\n")):
        if idx:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        add_markdown_runs(paragraph, line)


def add_heading(document: Document, text: str, style: str, page_break: bool) -> None:
    if page_break and document.paragraphs:
        document.add_page_break()
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)


def is_scene_break(block: list[str]) -> bool:
    if len(block) != 1:
        return False
    normalized = block[0].replace(" ", "")
    return normalized in {"***", "---"}


def build_document(markdown_text: str) -> Document:
    document = Document()
    configure_page(document)
    configure_styles(document)

    blocks = split_blocks(markdown_text)
    seen_title = False
    first_h3_seen = False
    after_display_block = False

    for block in blocks:
        if len(block) == 1:
            heading_match = HEADING_RE.match(block[0])
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()

                if level == 1:
                    if not seen_title:
                        add_heading(document, text, "PrintSourceTitle", page_break=False)
                        seen_title = True
                    else:
                        add_heading(
                            document,
                            text,
                            "PrintSourceChapterHeading",
                            page_break=True,
                        )
                    after_display_block = True
                    continue

                if level == 2:
                    add_heading(
                        document, text, "PrintSourceSubheading", page_break=True
                    )
                    after_display_block = True
                    continue

                if level == 3:
                    page_break = first_h3_seen
                    style = "PrintSourceDedication" if page_break else "PrintSourceByline"
                    add_heading(document, text, style, page_break=page_break)
                    first_h3_seen = True
                    after_display_block = True
                    continue

        if is_scene_break(block):
            add_markdown_paragraph(document, "* * *", "PrintSourceSceneBreak")
            after_display_block = True
            continue

        paragraph_text = "\n".join(block)
        paragraph_style = (
            "PrintSourceBodyFirst" if after_display_block else "PrintSourceBody"
        )

        if "Copyright" in paragraph_text or paragraph_text.startswith("Paperback ISBN:"):
            paragraph_style = "PrintSourceCopyright"

        add_markdown_paragraph(document, paragraph_text, paragraph_style)
        after_display_block = False

    return document


def main() -> int:
    args = parse_args()
    input_path = Path(args.input)
    output_path = Path(args.output)

    markdown_text = input_path.read_text(encoding="utf-8")
    document = build_document(markdown_text)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    print(f"Built print-source DOCX at: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
